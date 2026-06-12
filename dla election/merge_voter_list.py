import os
import sys
import docx
import pandas as pd
import json
import difflib
import re

# Ensure standard output uses UTF-8 to handle Nepali Unicode characters
sys.stdout.reconfigure(encoding='utf-8')

# Define file paths
docx_path = r"c:\xampp\htdocs\codes\election\dla election\final-voter-list.docx"
xlsx_path = r"c:\xampp\htdocs\codes\election\dla election\phone-numbers.xlsx"
json_path = r"c:\xampp\htdocs\codes\election\dla election\address.json"
output_path = r"c:\xampp\htdocs\codes\election\dla election\final-voter-list-merged.xlsx"

# 1. Initialize npttf2utf FontMapper
import npttf2utf
map_json_path = os.path.join(os.path.dirname(npttf2utf.__file__), 'map.json')
fm = npttf2utf.FontMapper(map_json_path)

def clean_nepali_name(name):
    if not name:
        return ""
    name = name.strip()
    # Remove parenthetical text (e.g. "(सुवास)", "(खाती)", "(कपिल)")
    name = re.sub(r'\(.*?\)', '', name)
    
    # Normalize spelling variations and prefixes globally
    name = name.replace("बहादूर", "बहादुर").replace("वहादूर", "बहादुर").replace("वहादुर", "बहादुर")
    name = name.replace("र्प्साद", "प्रसाद").replace("कमार", "कुमार")
    name = name.replace("डा.", "").replace("डा ", "").replace("एड.", "").replace("एड ", "")
    name = name.replace("ङ्ग", "ं") # normalize nasal conjunct to anusvara dot
    
    words = name.split()
    if not words:
        return ""
        
    stopwords = ["प्रसाद", "बहादुर", "कुमार", "कुमारी", "शर्मा", "दत्त", "राज", "लाल", "शरण", "नाथ", "प्र", "सिनियर", "कु"]
    compound_stopwords = ["प्रसाद", "बहादुर", "कुमार", "कुमारी", "कु"]
    
    cleaned_words = []
    for i, w in enumerate(words):
        w_clean = w.replace(".", "").replace(",", "")
        
        # If it is a stopword and not the last word, skip it
        if w_clean in stopwords and i < len(words) - 1:
            continue
            
        # Clean compound middle names from non-last words (only using compound_stopwords)
        if i < len(words) - 1:
            for sw in compound_stopwords:
                if sw in w:
                    w = w.replace(sw, "")
        cleaned_words.append(w)
        
    name_clean = "".join(cleaned_words)
    for char in [" ", ".", ",", "-", "“", "”", '"', "'", "(", ")", "/", "÷"]:
        name_clean = name_clean.replace(char, "")
    return name_clean

def is_unicode(text):
    if not text:
        return False
    # Nepali Devanagari range is 0x0900 to 0x097F
    return any(0x0900 <= ord(c) <= 0x097F for c in text)

def convert_if_preeti(text):
    if not text:
        return ""
    text_str = str(text).strip()
    if is_unicode(text_str):
        return text_str
    try:
        # Convert using FontMapper
        return fm.map_to_unicode(text_str)
    except Exception:
        return text_str

def clean_license(val):
    if pd.isna(val) or val is None:
        return ""
    val_str = str(val).strip()
    if val_str.endswith('.0'):
        val_str = val_str[:-2]
    return val_str

# 2. Load JSON database
print("Loading address.json...")
with open(json_path, 'r', encoding='utf-8') as f:
    json_data = json.load(f)

json_by_license = {}
json_by_name = {}
for item in json_data:
    if len(item) > 0:
        lic = clean_license(item[0])
        if lic:
            json_by_license.setdefault(lic, []).append(item)
        if len(item) > 1:
            clean_name = clean_nepali_name(item[1])
            if clean_name:
                json_by_name.setdefault(clean_name, []).append(item)

# 3. Load XLSX phone number database
print("Loading phone-numbers.xlsx...")
df_excel = pd.read_excel(xlsx_path)
df_clean = df_excel.iloc[7:].copy()
df_clean.columns = [df_excel.iloc[6, col_idx] for col_idx in range(df_excel.shape[1])]

xlsx_by_license = {}
xlsx_by_name = {}
all_xlsx_records = []
for idx, row in df_clean.iterrows():
    lic = clean_license(row['k|=k=g+='])
    row_dict = row.to_dict()
    row_dict['cleaned_lic'] = lic
    row_dict['name_unicode'] = convert_if_preeti(row['gfd'])
    row_dict['phone_clean'] = str(row[';Dks{ g+=']).strip() if pd.notna(row[';Dks{ g+=']) else ""
    
    if lic:
        xlsx_by_license.setdefault(lic, []).append(row_dict)
    clean_name = clean_nepali_name(row_dict['name_unicode'])
    if clean_name:
        xlsx_by_name.setdefault(clean_name, []).append(row_dict)
    all_xlsx_records.append(row_dict)

# 4. Read DOCX voter list
print("Reading final-voter-list.docx...")
doc = docx.Document(docx_path)

# Counters for reporting
total_records = 0
matched_address_by_lic = 0
matched_address_by_name = 0
matched_address_by_fuzzy_name = 0
matched_phone_by_lic = 0
matched_phone_by_name = 0
matched_phone_by_fuzzy_name = 0

merged_rows = []

def process_table(table, table_name, is_life_members=False):
    global total_records, matched_address_by_lic, matched_address_by_name, matched_address_by_fuzzy_name
    global matched_phone_by_lic, matched_phone_by_name, matched_phone_by_fuzzy_name
    
    print(f"\nProcessing {table_name}...")
    
    for idx, row in enumerate(table.rows[1:]): # Skip header
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) <= 1:
            continue
            
        lic_raw = cells[1]
        name_preeti = cells[2]
        title_raw = cells[3] if len(cells) > 3 else ""
        
        # Clean license and convert name
        lic = clean_license(lic_raw)
        name_unicode = convert_if_preeti(name_preeti)
        
        if not lic and not name_preeti:
            continue
            
        total_records += 1
        
        # Determine Renewal Status
        if is_life_members:
            # Life members don't have a renewal status column, leave it blank as planned
            renewal_status = ""
        else:
            # Table 0: column 5 is renewal status
            renewal_status = cells[5] if len(cells) > 5 else ""
            
        # Match Name and Address in JSON
        matched_name = name_unicode
        matched_address = ""
        address_match_type = "None"
        
        # A. Look up by license in JSON
        json_candidates = json_by_license.get(lic, [])
        best_json_match = None
        best_json_ratio = 0.0
        
        for cand in json_candidates:
            cand_name = cand[1]
            ratio = difflib.SequenceMatcher(None, clean_nepali_name(cand_name), clean_nepali_name(name_unicode)).ratio()
            if ratio > best_json_ratio:
                best_json_ratio = ratio
                best_json_match = cand
                
        if best_json_match and best_json_ratio >= 0.85:
            matched_name = best_json_match[1] # official name from registry
            matched_address = best_json_match[2]
            address_match_type = "License"
            matched_address_by_lic += 1
        else:
            # B. Fallback: Lookup by exact cleaned name in JSON (O(1))
            cleaned_target_name = clean_nepali_name(name_unicode)
            exact_name_candidates = json_by_name.get(cleaned_target_name, [])
            if exact_name_candidates:
                # Use the first exact match
                matched_name = exact_name_candidates[0][1]
                matched_address = exact_name_candidates[0][2]
                address_match_type = "Exact Name"
                matched_address_by_name += 1
            else:
                # C. Fallback: Fuzzy match only among candidates starting with same first character (O(N_filtered))
                best_fallback_match = None
                best_fallback_ratio = 0.0
                first_char = name_unicode[0] if name_unicode else ""
                
                # Filter items starting with same character to keep it extremely fast
                for item in json_data:
                    if len(item) > 1 and item[1] and item[1][0] == first_char:
                        cand_name = item[1]
                        ratio = difflib.SequenceMatcher(None, clean_nepali_name(cand_name), cleaned_target_name).ratio()
                        if ratio > best_fallback_ratio:
                            best_fallback_ratio = ratio
                            best_fallback_match = item
                            
                if best_fallback_match and best_fallback_ratio >= 0.85:
                    matched_name = best_fallback_match[1]
                    matched_address = best_fallback_match[2]
                    address_match_type = "Fuzzy Name Fallback"
                    matched_address_by_fuzzy_name += 1
                
        # Match Phone Number in XLSX
        phone_number = ""
        phone_match_type = "None"
        
        # A. Look up by license in XLSX
        xlsx_candidates = xlsx_by_license.get(lic, [])
        best_xlsx_match = None
        best_xlsx_ratio = 0.0
        
        for cand in xlsx_candidates:
            cand_name = cand['name_unicode']
            ratio = difflib.SequenceMatcher(None, clean_nepali_name(cand_name), clean_nepali_name(name_unicode)).ratio()
            if ratio > best_xlsx_ratio:
                best_xlsx_ratio = ratio
                best_xlsx_match = cand
                
        if best_xlsx_match and best_xlsx_ratio >= 0.85:
            # If the best match row has a phone number, use it
            if best_xlsx_match['phone_clean']:
                phone_number = best_xlsx_match['phone_clean']
            else:
                # If chosen candidate has empty phone, see if any other candidate for this license has a phone and matches name
                for cand in xlsx_candidates:
                    cand_name = cand['name_unicode']
                    cand_ratio = difflib.SequenceMatcher(None, clean_nepali_name(cand_name), clean_nepali_name(name_unicode)).ratio()
                    if cand['phone_clean'] and cand_ratio >= 0.85:
                        phone_number = cand['phone_clean']
                        break
            phone_match_type = "License"
            matched_phone_by_lic += 1
        else:
            # B. Fallback: Lookup by exact cleaned name in XLSX (O(1))
            cleaned_target_name = clean_nepali_name(name_unicode)
            exact_xlsx_candidates = xlsx_by_name.get(cleaned_target_name, [])
            # Filter ones that have a valid phone number
            exact_with_phone = [c for c in exact_xlsx_candidates if c['phone_clean']]
            if exact_with_phone:
                phone_number = exact_with_phone[0]['phone_clean']
                phone_match_type = "Exact Name"
                matched_phone_by_name += 1
            elif exact_xlsx_candidates:
                phone_number = exact_xlsx_candidates[0]['phone_clean']
                phone_match_type = "Exact Name"
                matched_phone_by_name += 1
            else:
                # C. Fallback: Fuzzy match only among candidates starting with same first character (O(N_filtered))
                best_fallback_xlsx = None
                best_fallback_ratio = 0.0
                first_char = name_unicode[0] if name_unicode else ""
                
                for rec in all_xlsx_records:
                    cand_name = rec['name_unicode']
                    if cand_name and cand_name[0] == first_char:
                        ratio = difflib.SequenceMatcher(None, clean_nepali_name(cand_name), cleaned_target_name).ratio()
                        if ratio > best_fallback_ratio:
                            best_fallback_ratio = ratio
                            best_fallback_xlsx = rec
                            
                if best_fallback_xlsx and best_fallback_ratio >= 0.85:
                    phone_number = best_fallback_xlsx['phone_clean']
                    phone_match_type = "Fuzzy Name Fallback"
                    matched_phone_by_fuzzy_name += 1
                
        # Format phone number (remove trailing .0 or float artifacts)
        if phone_number.endswith('.0'):
            phone_number = phone_number[:-2]
        if phone_number == "nan" or phone_number == "None":
            phone_number = ""
            
        merged_rows.append({
            "S.N": len(merged_rows) + 1,
            "Name": matched_name,
            "License": lic,
            "Renewal Status": renewal_status,
            "Phone Number": phone_number,
            "Address": matched_address,
            "Note": "",
            "_SourceTable": table_name,
            "_AddressMatchType": address_match_type,
            "_PhoneMatchType": phone_match_type,
            "_DocxNameUnicode": name_unicode
        })

# Process Table 0 (Ordinary Members)
process_table(doc.tables[0], "Table 0 (Ordinary)", is_life_members=False)

# Process Table 1 (Life Members)
process_table(doc.tables[1], "Table 1 (Life Members)", is_life_members=True)

# Create DataFrame
df_out = pd.DataFrame(merged_rows)

# Create clean output DataFrame with only the requested columns
df_clean_out = df_out[["S.N", "Name", "License", "Renewal Status", "Phone Number", "Address", "Note"]].copy()

# Save to Excel
df_clean_out.to_excel(output_path, index=False)
print(f"\nSaved final merged list to: {output_path}")

# Display Summary Statistics
print("\n=== SUMMARY STATISTICS ===")
print(f"Total DOCX rows processed: {total_records}")
print(f"Addresses matched by License: {matched_address_by_lic} ({matched_address_by_lic/total_records*100:.2f}%)")
print(f"Addresses matched by Exact Name: {matched_address_by_name} ({matched_address_by_name/total_records*100:.2f}%)")
print(f"Addresses matched by Fuzzy Name: {matched_address_by_fuzzy_name} ({matched_address_by_fuzzy_name/total_records*100:.2f}%)")
print(f"Total Addresses matched: {matched_address_by_lic + matched_address_by_name + matched_address_by_fuzzy_name} ({(matched_address_by_lic + matched_address_by_name + matched_address_by_fuzzy_name)/total_records*100:.2f}%)")
print(f"Phones matched by License: {matched_phone_by_lic} ({matched_phone_by_lic/total_records*100:.2f}%)")
print(f"Phones matched by Exact Name: {matched_phone_by_name} ({matched_phone_by_name/total_records*100:.2f}%)")
print(f"Phones matched by Fuzzy Name: {matched_phone_by_fuzzy_name} ({matched_phone_by_fuzzy_name/total_records*100:.2f}%)")
print(f"Total Phones matched: {matched_phone_by_lic + matched_phone_by_name + matched_phone_by_fuzzy_name} ({(matched_phone_by_lic + matched_phone_by_name + matched_phone_by_fuzzy_name)/total_records*100:.2f}%)")

# Highlight conflicting license numbers and how they resolved
print("\n=== VERIFYING CONFLICT RESOLUTIONS ===")
conflict_licenses = ['3872', '6749', '6824', '7964', '18953']
df_conflicts = df_out[df_out['License'].isin(conflict_licenses)]
print(df_conflicts[["S.N", "License", "Name", "Phone Number", "Address", "_SourceTable", "_AddressMatchType", "_PhoneMatchType"]].to_string())

# Check for rows with completely missing phone and address
missing_both = df_out[(df_out['Phone Number'] == '') & (df_out['Address'] == '')]
print(f"\nRows missing BOTH Phone and Address: {len(missing_both)}")
if len(missing_both) > 0:
    print("Sample missing rows:")
    print(missing_both[["S.N", "License", "Name", "_SourceTable"]].head(15).to_string())
